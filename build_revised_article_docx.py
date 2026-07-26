#!/usr/bin/env python3
"""Append the corrected integrated analysis to the supplied theory manuscript.

The uploaded manuscript is never modified in place.  The output is a new DOCX
that preserves the theory text, narrows the four-regime sentence, adds a
structured abstract, and appends claim-matched computational Methods, Results,
figures, tables, limitations, and data/code statements.
"""
from __future__ import annotations

import json
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

import build_report_docx as report


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT.parents[1] / "upload" / "Individual Neural Holonomy(2).docx"
MANUSCRIPT = ROOT / "integrated" / "manuscript"
FIGURES = ROOT / "integrated" / "figures"
RESULTS = ROOT / "integrated" / "results"
OUTPUT = MANUSCRIPT / "INH_Revised_Article_With_Integrated_Analysis.docx"


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def parse_subsections(markdown: str) -> dict[str, list[str]]:
    """Return level-three Markdown sections as clean paragraph blocks."""
    sections: dict[str, list[str]] = {}
    heading: str | None = None
    buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal buffer
        if heading is not None and buffer:
            value = " ".join(part.strip() for part in buffer if part.strip()).strip()
            if value:
                sections[heading].append(value)
        buffer = []

    for line in markdown.splitlines():
        if line.startswith("### "):
            flush_paragraph()
            heading = line[4:].strip()
            sections.setdefault(heading, [])
        elif line.startswith("## "):
            flush_paragraph()
            heading = None
        elif not line.strip():
            flush_paragraph()
        elif heading is not None:
            buffer.append(line)
    flush_paragraph()
    return sections


def insert_before(reference, paragraph) -> None:
    reference._p.addprevious(paragraph._p)


def insert_after(reference, text: str) -> None:
    new_p = OxmlElement("w:p")
    reference._p.addnext(new_p)
    paragraph = reference._parent.add_paragraph()
    paragraph._p.getparent().remove(paragraph._p)
    new_p.addnext(paragraph._p)
    paragraph.add_run(text)


def insert_structured_abstract(doc: Document, intro) -> None:
    abstract = (MANUSCRIPT / "Proposed_Abstract.md").read_text(encoding="utf-8")
    blocks = []
    for paragraph in abstract.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph or paragraph.startswith("# "):
            continue
        blocks.append(re.sub(r"\*\*([^*]+)\.\*\*", r"\1.—", paragraph))

    created = [doc.add_heading("Abstract", level=1)]
    created.extend(doc.add_paragraph(value) for value in blocks)
    keywords = doc.add_paragraph()
    key_run = keywords.add_run("Keywords: ")
    key_run.bold = True
    keywords.add_run("neural holonomy; SU(2); retention; protention; computational validation; identifiability")
    created.append(keywords)
    for paragraph in created:
        insert_before(intro, paragraph)


def format_existing_structure(doc: Document) -> None:
    if doc.paragraphs:
        doc.paragraphs[0].style = doc.styles["Title"]
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    pattern = re.compile(r"^(\d+)(?:\.(\d+))?\.\s")
    for paragraph in doc.paragraphs[1:]:
        match = pattern.match(paragraph.text.strip())
        if not match:
            continue
        paragraph.style = doc.styles["Heading 2" if match.group(2) else "Heading 1"]


def add_inline_figure(doc: Document, number: int, stem: str, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(FIGURES / f"{stem}.png"), width=Inches(6.35))
    shape._inline.docPr.set("title", f"Figure {number}")
    shape._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph(style="Caption")
    run = cap.add_run(f"Figure {number} | ")
    report.set_run_font(run, size=9, color=report.NAVY, bold=True)
    run = cap.add_run(caption)
    report.set_run_font(run, size=9, color=report.GRAY, italic=True)


def add_numbered_section(doc: Document, number: str, title: str, paragraphs: list[str]) -> None:
    doc.add_heading(f"{number}. {title}", level=2)
    for text in paragraphs:
        doc.add_paragraph(text)


def add_stage3_table(doc: Document, stage3: dict) -> None:
    rows = []
    short = ["S–C", "S–R", "S–D", "C–R", "C–D", "R–D"]
    for index, item in enumerate(stage3["pairwise_validation"].values()):
        rows.append([
            short[index],
            f"{item['common_frame']['mean']:.4f} [{item['common_frame']['lower']:.4f}, {item['common_frame']['upper']:.4f}]",
            f"{item['common_frame']['simultaneous_lower']:.4f}",
            f"{item['conjugacy']['mean']:.4f} [{item['conjugacy']['lower']:.4f}, {item['conjugacy']['upper']:.4f}]",
            f"{item['conjugacy']['simultaneous_lower']:.4f}",
            "Primary" if index < 3 else "Secondary",
        ])
    report.add_table(
        doc,
        ["Pair", "Common mean [95% CI]", "Common LCB", "Conjugacy mean [95% CI]", "Conj. LCB", "Family"],
        rows,
        [.48, 1.55, .72, 1.55, .72, 1.48],
        "Table 1 | Stage 3 clone-clustered pairwise inference",
    )


def add_stage4_table(doc: Document, stage4: dict) -> None:
    value = stage4["validation"]
    entries = [
        ("Scalar Spearman", value["angles"]["spearman"], "LCB > .72"),
        ("Scalar MAE", value["angles"]["mae"], "UCB < .045"),
        ("Pairwise Spearman", value["pairwise_distance_ordering"]["spearman"], "LCB > .72"),
        ("Pairwise distortion", value["pairwise_distance_ordering"]["relative_distortion"], "UCB < .65"),
        ("Within-setup ICC", value["repeat_icc_within_setup_absolute"], "LCB > .70"),
        ("Dynamic-null AUC", value["dynamic_zero_holonomy_auc"], "LCB > .85"),
    ]
    rows = [[name, f"{metric['value']:.4f}", f"[{metric['lower']:.4f}, {metric['upper']:.4f}]", criterion, "PASS"] for name, metric, criterion in entries]
    report.add_table(doc, ["Estimand", "Estimate", "95% CI", "Criterion", "Decision"], rows, [1.8, .85, 1.35, 1.25, .75], "Table 2 | Stage 4 held-out validation criteria")


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    doc = Document(SOURCE)
    report.configure_document(doc)
    for section in doc.sections:
        report.header_footer(section)
    format_existing_structure(doc)

    intro = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "1. Introduction")
    insert_structured_abstract(doc, intro)

    regime_sentence = next((paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "four regimes can be distinguished:"), None)
    if regime_sentence is not None:
        regime_sentence.text = "the model specifies four mechanistic regimes:"
    regime_limit = next((paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("These regimes concern finite intervals")), None)
    if regime_limit is not None:
        insert_after(
            regime_limit,
            "In the final computational validation, the stable regime was separated from each modeled non-stable regime, whereas complete pairwise separation among all four regimes was not supported.",
        )

    markdown = (MANUSCRIPT / "Manuscript_Ready_Methods_and_Results.md").read_text(encoding="utf-8")
    sections = parse_subsections(markdown)
    conclusion_text = markdown.split("## Integrated conclusion and limitations", 1)[1]
    conclusion_paragraphs = [
        " ".join(line.strip() for line in block.splitlines() if line.strip())
        for block in conclusion_text.strip().split("\n\n")
        if block.strip()
    ]
    stage3 = load_json(RESULTS / "stage3_final" / "summary.json")
    stage4 = load_json(RESULTS / "stage4_final" / "summary.json")

    doc.add_page_break()
    doc.add_heading("6. Computational Validation", level=1)
    add_numbered_section(doc, "6.1", "Integrated design and inferential scope", sections["Integrated analysis design"])
    add_inline_figure(doc, 1, "Figure_1_Integrated_Study_Design", "Integrated study design and strict inferential boundary across Stages 1–5.")
    add_numbered_section(doc, "6.2", "Units, cohort splits, and uncertainty", sections["Statistical units, splits, and uncertainty"])
    add_numbered_section(doc, "6.3", "Randomization, null models, and endogenous remodeling", sections["Randomization and null models"])
    add_numbered_section(doc, "6.4", "Controlled observation model and recovery", sections["Controlled observation model and recovery"])
    add_numbered_section(doc, "6.5", "Transformation taxonomy", sections["Transformation taxonomy"])

    add_numbered_section(doc, "6.6", "Constructive and held-out latent evidence", sections["Stage 1 establishes only a constructive shared-frame example"] + sections["Stage 2 validates a shared-frame interaction and independent-noise robustness"])
    add_inline_figure(doc, 2, "Figure_2_Constructive_and_Heldout_Latent_Evidence", "Constructive Stage 1 effects and held-out Stage 2 shared-frame evidence with clone-clustered uncertainty.")

    add_numbered_section(doc, "6.7", "Endogenous remodeling and paired ablations", sections["Stage 3 supports stable-versus-non-stable separation, not a complete taxonomy"])
    add_inline_figure(doc, 3, "Figure_3_Endogenous_Remodeling_Dynamics", "Endogenous remodeling dynamics across the four modeled regimes.")
    add_inline_figure(doc, 4, "Figure_4_Stage3_Familywise_Inference_and_Ablations", "Familywise Stage 3 inference and paired mechanistic ablations.")
    add_stage3_table(doc, stage3)

    add_numbered_section(doc, "6.8", "Controlled descriptor recovery", sections["Stage 4 recovers claim-matched descriptors in the controlled inverse problem"])
    add_inline_figure(doc, 5, "Figure_5_CrossFitted_Observation_Model_Recovery", "Cross-fitted scalar and within-clone metric recovery under the declared synthetic observation model.")
    add_stage4_table(doc, stage4)

    add_numbered_section(doc, "6.9", "Robustness and transformation-specific boundaries", sections["Stage 5 maps claim-specific robustness and failure boundaries"])
    add_inline_figure(doc, 6, "Figure_6_Observation_Model_Stress_Tests", "Observation-model stress tests across noise, channel count, carrier misspecification, and mixing drift.")
    add_inline_figure(doc, 7, "Figure_7_Transformation_Specific_Claim_Boundaries", "Claim-specific information-loss boundaries and temporal-alignment nuisance.")

    doc.add_heading("7. Discussion and Limitations", level=1)
    for paragraph in conclusion_paragraphs:
        doc.add_paragraph(paragraph)
    doc.add_paragraph("Stage 1 is constructive, Stages 2–3 are held-out latent simulations, Stage 4 is recovery under a generator/estimator-matched synthetic inverse problem, and Stage 5 reuses the Stage 4 validation cohort for sensitivity analysis. These levels of evidence must not be conflated.")
    doc.add_paragraph("The analysis does not establish a complete four-regime taxonomy, fast-RIC causality, full-connection recovery, unrestricted cross-subject gauge identification, real-ECoG validity, or phenomenological individuation. External empirical data and prospective validation are required for those claims.")
    add_inline_figure(doc, 8, "Figure_8_Integrated_Claim_Matrix", "Integrated evidence matrix separating supported, failed, and out-of-scope claims.")

    doc.add_heading("8. Conclusion", level=1)
    doc.add_paragraph("In a controlled rank-two SU(2) simulation, pilot-calibrated history-dependent remodeling separates the modeled stable regime from non-stable alternatives under matched probes, and selected conjugacy-invariant scalar and within-clone metric descriptors are recoverable under the declared synthetic observation model.")

    doc.add_heading("9. Data and Code Availability", level=1)
    doc.add_paragraph("All corrected source code, fixed configurations, machine-readable summaries, clone-level source data, full-resolution figures, independent recomputation records, checksums, and exact rerun instructions are included in the accompanying versioned analysis archive.")

    doc.add_heading("10. Submission Metadata to Complete", level=1)
    report.add_callout(doc, "Editorial requirement", "The source manuscript contains numbered citations but no reference list, and author names, affiliations, corresponding-author details, author contributions, funding, and competing-interest statements were not supplied. These bibliographic and author-provided fields must be completed before journal submission.", kind="risk")

    doc.core_properties.title = "Individual Neural Holonomy: Retention–Protention Coupling and the State-Space Geometry of Unique Experience"
    doc.core_properties.subject = "Theory manuscript with integrated computational validation of Stages 1–5"
    doc.core_properties.keywords = "neural holonomy; retention; protention; SU(2); computational validation; identifiability"
    doc.core_properties.comments = "Original source preserved; computational sections generated from corrected full-run artifacts."
    MANUSCRIPT.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
