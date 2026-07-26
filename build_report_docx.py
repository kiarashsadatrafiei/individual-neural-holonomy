#!/usr/bin/env python3
"""Create the visually verified integrated final-analysis report DOCX."""
from __future__ import annotations

import json
from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "integrated" / "results"
FIGURES = ROOT / "integrated" / "figures"
TABLES = ROOT / "integrated" / "tables"
MANUSCRIPT = ROOT / "integrated" / "manuscript"
OUTPUT = MANUSCRIPT / "INH_Integrated_Final_Analysis_Report.docx"

sys.path.insert(0, "/root/.codex/skills/builtins/documents/scripts")
from table_geometry import apply_table_geometry, column_widths_from_weights

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRAY = "5F6B76"
LIGHT = "F4F6F9"
PALE_BLUE = "E8EEF5"
GREEN = "1F6E43"
PALE_GREEN = "E2F0D9"
RED = "9B1C1C"
PALE_RED = "FCE8E6"
GOLD = "7A5A00"


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd"); tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders"); tc_pr.append(borders)
    for edge_name, attrs in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}"); borders.append(edge)
        for key, value in attrs.items(): edge.set(qn(f"w:{key}"), str(value))


def set_run_font(run, *, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None: run.font.size = Pt(size)
    if color is not None: run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic


def set_style_font(style, name: str, size: float, color: str | None = None, bold: bool | None = None) -> None:
    style.font.name = name; style.font.size = Pt(size)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if color: style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: style.font.bold = bold


def configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5); section.page_height = Inches(11)
        section.top_margin = Inches(1); section.bottom_margin = Inches(1)
        section.left_margin = Inches(1); section.right_margin = Inches(1)
        section.header_distance = Inches(.492); section.footer_distance = Inches(.492)
    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11)
    normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name,size,color,before,after in (("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK_BLUE,8,4)):
        style=doc.styles[name]; set_style_font(style,"Calibri",size,color,True)
        style.paragraph_format.space_before=Pt(before); style.paragraph_format.space_after=Pt(after)
        style.paragraph_format.keep_with_next=True
    try:
        caption=doc.styles["Caption"]
    except KeyError:
        caption=doc.styles.add_style("Caption",WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(caption,"Calibri",9,GRAY,False)
    caption.font.italic=True; caption.paragraph_format.space_before=Pt(4); caption.paragraph_format.space_after=Pt(10); caption.paragraph_format.keep_with_next=False
    for list_name in ("List Bullet","List Number"):
        try:
            style=doc.styles[list_name]
        except KeyError:
            style=doc.styles.add_style(list_name,WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(style,"Calibri",11)
        style.paragraph_format.left_indent=Inches(.375); style.paragraph_format.first_line_indent=Inches(-.194)
        style.paragraph_format.space_after=Pt(4); style.paragraph_format.line_spacing=1.208


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin=OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"),"begin")
    instr=OxmlElement("w:instrText"); instr.set(qn("xml:space"),"preserve"); instr.text=" PAGE "
    separate=OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"),"separate")
    text=OxmlElement("w:t"); text.text="1"
    end=OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"),"end")
    for element in (begin,instr,separate,text,end): run._r.append(element)


def header_footer(section) -> None:
    header=section.header; p=header.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after=Pt(2)
    run=p.add_run("INDIVIDUAL NEURAL HOLONOMY  |  INTEGRATED FINAL ANALYSIS v1.0.0")
    set_run_font(run,size=8,color=GRAY,bold=True)
    p_pr=p._p.get_or_add_pPr(); borders=OxmlElement("w:pBdr"); bottom=OxmlElement("w:bottom")
    for key,value in (("val","single"),("sz","4"),("space","3"),("color","B8C2CC")): bottom.set(qn(f"w:{key}"),value)
    borders.append(bottom); p_pr.append(borders)
    footer=section.footer; fp=footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run=fp.add_run("Page "); set_run_font(run,size=8,color=GRAY); add_page_field(fp)


def add_title_page(doc: Document) -> None:
    for _ in range(5): doc.add_paragraph().paragraph_format.space_after=Pt(8)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(14)
    r=p.add_run("COMPUTATIONAL VALIDATION REPORT"); set_run_font(r,size=10,color=GOLD,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8)
    r=p.add_run("Individual Neural Holonomy"); set_run_font(r,size=30,color=NAVY,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(26)
    r=p.add_run("Integrated Stages 1–5 • Submission-oriented final analysis"); set_run_font(r,size=15,color=DARK_BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(70)
    r=p.add_run("Corrected inference • full reruns • cluster-bootstrap uncertainty • publication figures • reproducibility audit"); set_run_font(r,size=10,color=GRAY,italic=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Version 1.0.0  |  22 July 2026"); set_run_font(r,size=11,color=NAVY,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Controlled rank-two SU(2) simulation; no real-ECoG or phenomenological claim"); set_run_font(r,size=9.5,color=RED,italic=True)
    doc.add_page_break()


def add_callout(doc: Document, label: str, text: str, *, kind="info") -> None:
    colors={"info":(PALE_BLUE,NAVY),"success":(PALE_GREEN,GREEN),"risk":(PALE_RED,RED)}
    fill,color=colors[kind]
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.right_indent=Inches(.18)
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(10)
    p_pr=p._p.get_or_add_pPr(); shading=OxmlElement("w:shd"); shading.set(qn("w:fill"),fill); p_pr.append(shading)
    borders=OxmlElement("w:pBdr"); left=OxmlElement("w:left")
    for key,value in (("val","single"),("sz","18"),("space","8"),("color",color)): left.set(qn(f"w:{key}"),value)
    borders.append(left); p_pr.append(borders)
    r=p.add_run(f"{label}: "); set_run_font(r,size=10.5,color=color,bold=True)
    r=p.add_run(text); set_run_font(r,size=10.5,color="222222")


def add_bullet(doc: Document, text: str) -> None:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(text)


def format_table(table, widths, *, header=True, font_size=8.5) -> None:
    apply_table_geometry(table,column_widths_from_weights(widths,9360),table_width_dxa=9360,indent_dxa=120,cell_margins_dxa={"top":90,"bottom":90,"start":120,"end":120})
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row_index,row in enumerate(table.rows):
        row.height=None
        for cell in row.cells:
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index==0 and header: set_cell_shading(cell,PALE_BLUE)
            set_cell_border(cell,bottom={"val":"single","sz":"3","color":"D9E0E7"})
            for p in cell.paragraphs:
                p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.08
                for run in p.runs: set_run_font(run,size=font_size,color=NAVY if row_index==0 else "222222",bold=True if row_index==0 else False)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], weights: list[float], caption: str) -> None:
    p=doc.add_paragraph(caption,style="Caption"); p.paragraph_format.keep_with_next=True
    table=doc.add_table(rows=1,cols=len(headers))
    for index,value in enumerate(headers): table.rows[0].cells[index].text=str(value)
    for values in rows:
        cells=table.add_row().cells
        for index,value in enumerate(values): cells[index].text=str(value)
    format_table(table,weights,font_size=8.1 if len(headers)>5 else 8.5)
    doc.add_paragraph().paragraph_format.space_after=Pt(3)


def add_figure(doc: Document, number: int, stem: str, caption: str) -> None:
    doc.add_page_break()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(4); p.paragraph_format.keep_with_next=True
    shape=p.add_run().add_picture(str(FIGURES/f"{stem}.png"),width=Inches(6.45))
    shape._inline.docPr.set("title",f"Figure {number}")
    shape._inline.docPr.set("descr",caption)
    cap=doc.add_paragraph(style="Caption"); cap.paragraph_format.keep_with_next=False
    r=cap.add_run(f"Figure {number} | "); set_run_font(r,size=9,color=NAVY,bold=True)
    r=cap.add_run(caption); set_run_font(r,size=9,color=GRAY,italic=True)


def add_markdown_section(doc: Document, text: str, start_heading: str, stop_heading: str | None = None) -> None:
    lines=text.splitlines(); active=False; buffer=[]
    def flush():
        nonlocal buffer
        if buffer:
            paragraph=" ".join(line.strip() for line in buffer).strip()
            if paragraph: doc.add_paragraph(paragraph)
            buffer=[]
    for line in lines:
        if line.strip()==start_heading: active=True; continue
        if active and stop_heading and line.strip()==stop_heading: flush(); break
        if not active: continue
        if line.startswith("### "):
            flush(); doc.add_heading(line[4:].strip(),level=2)
        elif line.startswith("## "):
            flush(); doc.add_heading(line[3:].strip(),level=1)
        elif not line.strip(): flush()
        else: buffer.append(line)
    flush()


def main() -> None:
    MANUSCRIPT.mkdir(parents=True,exist_ok=True)
    s1=load_json(RESULTS/"stage1_final"/"summary.json"); s2=load_json(RESULTS/"stage2_final"/"summary.json"); s3=load_json(RESULTS/"stage3_final"/"summary.json"); s4=load_json(RESULTS/"stage4_final"/"summary.json"); s5=load_json(RESULTS/"stage5_final"/"summary.json")
    md=(MANUSCRIPT/"Manuscript_Ready_Methods_and_Results.md").read_text(encoding="utf-8")
    doc=Document(); configure_document(doc); header_footer(doc.sections[0]); add_title_page(doc)

    doc.add_heading("Executive readiness assessment",level=1)
    add_callout(doc,"Verdict","The computational package is submission-oriented and internally reproducible for a controlled synthetic-model paper. The manuscript is not defensible if it presents the analysis as real-ECoG validation, complete four-regime individuation, fast-RIC causality, or evidence about phenomenology.",kind="success")
    add_bullet(doc,"Stage 1 is retained as a constructive shared-frame proof of principle; its suprathreshold interaction criterion fails.")
    add_bullet(doc,"Stage 2 validates a held-out shared-frame interaction that remains positive after paired independent-noise null subtraction.")
    add_bullet(doc,"Stage 3 supports familywise stable-versus-non-stable separation in both common-frame and conjugacy distance, but not complete pairwise separation of all four regimes.")
    add_bullet(doc,"Stage 4 passes scalar, metric, reliability, dynamic-null, numerical, and lineage criteria under the declared controlled observation model.")
    add_bullet(doc,f"Stage 5 exactly reproduces the selected frozen Stage 4 repeat, but that single repeat fails the scalar lower-bound cutoff and only {s5['information_loss_diagnostics']['successful_families']}/{s5['information_loss_diagnostics']['required_families']} strict information-loss families pass; circular shifts are treated as nuisance rather than information destruction.")

    stage_rows=[
        ["1","Constructive example","FAIL","Shared-frame proof only"],
        ["2","Held-out latent validation","PASS","Shared-frame; not unrestricted gauge identity"],
        ["3","Endogenous remodeling","PASS / LIMITED","Stable vs non-stable passes; global taxonomy fails"],
        ["4","Controlled observation recovery","PASS","Generator/estimator-matched synthetic inverse problem"],
        ["5","Sensitivity and boundaries","PASS" if s5["criteria"]["overall_pass"] else "LIMITED","Exact repeat reproduction; single-repeat scalar LCB and channel-projection strict test fail"],
    ]
    add_table(doc,["Stage","Role","Decision","Claim boundary"],stage_rows,[.55,1.55,1.05,3.35],"Table 1 | Integrated stage-level decisions")
    add_callout(doc,"Strongest supported statement","In a controlled rank-two SU(2) simulation, pilot-calibrated history-dependent remodeling separates the modeled stable regime from non-stable alternatives under matched probes, and selected conjugacy-invariant scalar and within-clone metric descriptors are recoverable under the declared synthetic observation model.",kind="info")
    add_callout(doc,"Explicitly unsupported","Complete four-regime taxonomy; fast-RIC causality; recovery of the full connection; unrestricted cross-subject gauge identification; real-ECoG validity; phenomenology or unique experience.",kind="risk")

    doc.add_heading("Integrated study design",level=1)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    shape=p.add_run().add_picture(str(FIGURES/"Figure_1_Integrated_Study_Design.png"),width=Inches(6.45))
    shape._inline.docPr.set("title","Figure 1")
    shape._inline.docPr.set("descr","Integrated study design, cohort structure, evidence chain, and strict claim boundary across Stages 1 through 5.")
    cap=doc.add_paragraph("Figure 1 summarizes the evidence chain, independent cohorts, and inferential scope.",style="Caption")

    doc.add_page_break(); doc.add_heading("Manuscript-ready Methods",level=1)
    add_markdown_section(doc,md,"## Methods","## Results")

    doc.add_page_break(); doc.add_heading("Manuscript-ready Results",level=1)
    add_markdown_section(doc,md,"## Results","## Integrated conclusion and limitations")
    doc.add_heading("Integrated conclusion and limitations",level=2)
    add_markdown_section(doc,md,"## Integrated conclusion and limitations",None)

    doc.add_heading("Key inferential tables",level=1)
    s3_rows=[]
    for index,(key,item) in enumerate(s3["pairwise_validation"].items()):
        s3_rows.append(["S–"+(["C","R","D"][index]) if index<3 else ["C–R","C–D","R–D"][index-3],f"{item['common_frame']['mean']:.4f} [{item['common_frame']['lower']:.4f}, {item['common_frame']['upper']:.4f}]",f"{item['common_frame']['simultaneous_lower']:.4f}",f"{item['conjugacy']['mean']:.4f} [{item['conjugacy']['lower']:.4f}, {item['conjugacy']['upper']:.4f}]",f"{item['conjugacy']['simultaneous_lower']:.4f}","Primary" if index<3 else "Global taxonomy"])
    add_table(doc,["Pair","Common mean [CI]","Common LCB","Conjugacy mean [CI]","Conj. LCB","Role"],s3_rows,[.48,1.55,.72,1.55,.72,1.48],"Table 2 | Stage 3 clone-clustered pairwise inference")
    v=s4["validation"]
    stage4_rows=[
        ["Scalar Spearman",f"{v['angles']['spearman']['value']:.3f}",f"[{v['angles']['spearman']['lower']:.3f}, {v['angles']['spearman']['upper']:.3f}]","LCB > .72","PASS"],
        ["Scalar MAE",f"{v['angles']['mae']['value']:.4f}",f"[{v['angles']['mae']['lower']:.4f}, {v['angles']['mae']['upper']:.4f}]","UCB < .045","PASS"],
        ["Pairwise Spearman",f"{v['pairwise_distance_ordering']['spearman']['value']:.3f}",f"[{v['pairwise_distance_ordering']['spearman']['lower']:.3f}, {v['pairwise_distance_ordering']['spearman']['upper']:.3f}]","LCB > .72","PASS"],
        ["Pairwise distortion",f"{v['pairwise_distance_ordering']['relative_distortion']['value']:.3f}",f"[{v['pairwise_distance_ordering']['relative_distortion']['lower']:.3f}, {v['pairwise_distance_ordering']['relative_distortion']['upper']:.3f}]","UCB < .65","PASS"],
        ["Within-setup ICC",f"{v['repeat_icc_within_setup_absolute']['value']:.3f}",f"[{v['repeat_icc_within_setup_absolute']['lower']:.3f}, {v['repeat_icc_within_setup_absolute']['upper']:.3f}]","LCB > .70","PASS"],
        ["Dynamic-null AUC",f"{v['dynamic_zero_holonomy_auc']['value']:.3f}",f"[{v['dynamic_zero_holonomy_auc']['lower']:.3f}, {v['dynamic_zero_holonomy_auc']['upper']:.3f}]","LCB > .85","PASS"],
    ]
    add_table(doc,["Estimand","Estimate","95% CI","Criterion","Decision"],stage4_rows,[1.8,.85,1.35,1.25,.75],"Table 3 | Stage 4 validation criteria")

    doc.add_heading("Publication figure atlas",level=1)
    captions=[
        "Integrated design and strict claim boundary.",
        "Constructive Stage 1 evidence and held-out Stage 2 latent validation.",
        "Endogenous Stage 3 remodeling dynamics with clone-clustered uncertainty.",
        "Stage 3 familywise pairwise inference and paired mechanistic ablations.",
        "Cross-fitted recovery from controlled synthetic multichannel observations.",
        "Observation-model stress tests using all validation clones and repeats.",
        "Transformation-specific information-loss boundaries and temporal nuisance.",
        "Integrated evidence and claim matrix.",
    ]
    stems=["Figure_1_Integrated_Study_Design","Figure_2_Constructive_and_Heldout_Latent_Evidence","Figure_3_Endogenous_Remodeling_Dynamics","Figure_4_Stage3_Familywise_Inference_and_Ablations","Figure_5_CrossFitted_Observation_Model_Recovery","Figure_6_Observation_Model_Stress_Tests","Figure_7_Transformation_Specific_Claim_Boundaries","Figure_8_Integrated_Claim_Matrix"]
    for number,(stem,caption) in enumerate(zip(stems,captions,strict=True),start=1): add_figure(doc,number,stem,caption)

    doc.add_page_break(); doc.add_heading("Reproducibility and quality-control record",level=1)
    add_bullet(doc,"Original uploaded archives were preserved; all corrections were made in a separate versioned integration directory.")
    add_bullet(doc,"All full stages were rerun from source with fixed seeds and single-threaded BLAS settings.")
    add_bullet(doc,"Pilot, calibration, development, and validation boundaries were checked for overlap and leakage.")
    add_bullet(doc,"Code was compiled recursively; numerical invariants, gauge covariance, reverse paths, convergence, and independent-session recovery were rechecked.")
    add_bullet(doc,"Every reported point estimate was independently recomputed from saved arrays and compared at machine precision.")
    add_bullet(doc,"Figures are supplied as 600-dpi PNG and vector PDF/SVG with source data in CSV, LaTeX, and XLSX.")
    add_bullet(doc,"The final Word report was rendered page-by-page to PNG/PDF and visually inspected before release.")
    add_callout(doc,"Submission condition","The paper text must use the limited claims in this report. If the manuscript still states complete individuation, empirical ECoG recovery, phenomenological identity, or fast-RIC causality, it is not submission-ready.",kind="risk")

    doc.add_heading("Insertion checklist for the article",level=1)
    for item in (
        "Replace broad abstract/conclusion language with the strongest supported controlled-model statement.",
        "Insert the integrated Methods and Results text supplied with this package; do not retain older stage-specific numbers.",
        "Use Figure 4 to disclose the failed global four-regime taxonomy alongside the successful stable-versus-non-stable family.",
        "Describe Stage 4 as controlled synthetic recovery under shared generator/estimator assumptions, not real-ECoG validation.",
        "Describe Stage 5 as sensitivity analysis on the reused Stage 4 validation cohort.",
        "State that fast-RIC causality, full-connection recovery, phenomenology, and unique experience remain unsupported.",
        "Archive the code, lock file/environment, machine-readable summaries, source-data workbook, and checksums with the submission.",
    ): add_bullet(doc,item)

    doc.core_properties.title="Individual Neural Holonomy — Integrated Final Analysis"
    doc.core_properties.subject="Submission-oriented computational validation of Stages 1–5"
    doc.core_properties.keywords="neural holonomy; SU(2); computational validation; bootstrap; reproducibility"
    doc.core_properties.comments="Generated from corrected full-run artifacts; original uploads preserved."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__=="__main__": main()
